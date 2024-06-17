from docx import Document
from docxcompose.composer import Composer
import os

doc0 = r'ProgramasPlantillas/docx/Manual_Etapas/Empty_Files/Manual_EmptyTemplatesE0.docx'
doc1 = r'ProgramasPlantillas/docx/Manual_Etapas/Empty_Files/Manual_EmptyTemplatesE1.docx'
doc2 = r'ProgramasPlantillas/docx/Manual_Etapas/Empty_Files/Manual_EmptyTemplatesE2.docx'
doc3 = r'ProgramasPlantillas/docx/Manual_Etapas/Empty_Files/Manual_EmptyTemplatesE3.docx'
doc4 = r'ProgramasPlantillas/docx/Manual_Etapas/Empty_Files/Manual_EmptyTemplatesE4.docx'
etapa0 = Document(doc0)
etapa1 = Document(doc1)
etapa2 = Document(doc2)
etapa3 = Document(doc3)
etapa4 = Document(doc4)
etapasTuple = (etapa1,etapa2,etapa3,etapa4)

composer = Composer(etapa0)
for e in etapasTuple:
    composer.append(e)

composer.save(os.path.join('ProgramasPlantillas/docx/Manual_Empty/','Manual_EmptyTemplates.docx'))





#print(os.getcwd())
#print(dirname)
#tablas=etapa0.tables
#print([tuple(c.text for c in r.cells) for r in tablas[0].rows])

#for row in tablas[len(tablas)-1].rows:
#    print(tuple(cell.text for cell in row.cells))
#    print("\n")

#for tables in tablas:
#    print(str(tables)+"\n")



# Check whether the specified 
# path exists or not 
#isExist = os.path.exists('ProgramasPlantillas/docx/Manual_FulfilledTemplatesV2.docx') 
#print(os.path.join(os.getcwd(),'\ProgramasPlantillas'))
#print(isExist)